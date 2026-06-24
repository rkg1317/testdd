"""Memo parsing must yield stable, citable passages."""

from docx import Document

from parser.memo_parser import parse_memo


def _make_docx(path):
    doc = Document()
    doc.add_heading("Review", level=1)
    doc.add_paragraph("Product Y in Region B declined in Q2.")
    doc.add_paragraph("")  # empty paragraph should be skipped
    doc.add_paragraph("A recovery is expected in Q3.")
    doc.save(path)
    return path


def test_passages_skip_empty_and_index_stably(tmp_path):
    p = _make_docx(tmp_path / "memo.docx")
    passages = parse_memo(p)

    # heading + two non-empty paragraphs = 3 passages (empty one skipped)
    assert len(passages) == 3
    assert passages[0].ref.kind == "heading"
    assert passages[1].content == "Product Y in Region B declined in Q2."
    assert passages[1].ref.document == "memo.docx"
    assert [pas.ref.passage for pas in passages] == [0, 1, 2]
    assert passages[2].to_dict()["citation"] == "memo.docx · passage 2"
