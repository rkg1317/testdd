"""Generate the synthetic, fully-generic fixtures for the cross-file showcase.

Everything here is invented and neutral (Region A/B/C, Product X/Y/Z) so the repo
can be pushed anywhere safely. The numbers are deterministic (no randomness) and
*engineered to tell a story* so the demo question has a clean, checkable answer:

    Q: "Across these regional files, which product's demand dropped most QoQ,
        and summarize the reason from the memo?"
    A: Product Y in Region B — Q1 avg 118 -> Q2 avg 68 (a 42% drop); the memo
       attributes it to a customer qualification delay.

Run:  python generate.py        (writes .xlsx + .docx into ./out/)
"""

from __future__ import annotations

from pathlib import Path

from openpyxl import Workbook
from docx import Document

OUT = Path(__file__).parent / "out"
MONTHS = ["Jan", "Feb", "Mar", "Apr", "May", "Jun"]  # H1: Q1 = Jan-Mar, Q2 = Apr-Jun
PRODUCTS = ["Product X", "Product Y", "Product Z"]

# Monthly demand per region. Region B / Product Y has the engineered Q2 cliff.
DEMAND = {
    "region_a": {
        "Product X": [100, 105, 110, 115, 120, 125],
        "Product Y": [80, 82, 84, 86, 88, 90],
        "Product Z": [60, 58, 56, 54, 52, 50],
    },
    "region_b": {
        "Product X": [90, 92, 94, 96, 98, 100],
        "Product Y": [120, 118, 116, 70, 68, 66],   # <-- sharp Q2 drop (the answer)
        "Product Z": [70, 71, 72, 73, 74, 75],
    },
    "region_c": {
        "Product X": [110, 109, 108, 107, 106, 105],
        "Product Y": [95, 96, 97, 98, 99, 100],
        "Product Z": [40, 42, 44, 46, 48, 50],
    },
}


def write_region_workbook(region_key: str) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Demand"
    ws.append(["Month"] + PRODUCTS)  # header row 1
    series = DEMAND[region_key]
    for i, month in enumerate(MONTHS):
        ws.append([month] + [series[p][i] for p in PRODUCTS])
    path = OUT / f"{region_key}.xlsx"
    wb.save(path)
    return path


def write_memo() -> Path:
    doc = Document()
    doc.add_heading("Quarterly Demand Review — Regions A/B/C", level=1)
    doc.add_paragraph(
        "This memo summarizes notable movements in regional product demand for the "
        "first half of the year. Figures are illustrative and used for demonstration."
    )
    doc.add_heading("Region B — Product Y", level=2)
    doc.add_paragraph(
        "In Q2, Product Y demand in Region B declined sharply versus Q1. The decline "
        "was driven by a customer qualification delay that pushed several committed "
        "orders beyond the quarter. A partial recovery is expected in Q3 once the "
        "outstanding qualifications complete."
    )
    doc.add_heading("Region A — Product X", level=2)
    doc.add_paragraph(
        "Product X in Region A continued a steady upward trend through H1, with no "
        "supply constraints reported."
    )
    doc.add_heading("Region C — General", level=2)
    doc.add_paragraph(
        "Region C demand was broadly stable across products, with a mild ramp in "
        "Product Z attributable to a new program qualification completing on schedule."
    )
    path = OUT / "memo.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    written = [write_region_workbook(k) for k in DEMAND] + [write_memo()]
    for p in written:
        print(f"wrote {p.relative_to(Path(__file__).parent)}")


if __name__ == "__main__":
    main()
